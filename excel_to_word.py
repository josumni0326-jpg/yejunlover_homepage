import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
import io
import os
from datetime import datetime
import threading

# Korean font for matplotlib
plt.rcParams['font.family'] = 'Malgun Gothic'
plt.rcParams['axes.unicode_minus'] = False


class ExcelToWordApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Excel → Word 보고서 생성기")
        self.root.geometry("620x520")
        self.root.resizable(False, False)
        self.root.configure(bg="#f0f2f5")

        self.excel_path = tk.StringVar()
        self.output_path = tk.StringVar()
        self.report_title = tk.StringVar(value="보고서")
        self.status = tk.StringVar(value="엑셀 파일을 선택해주세요.")
        self.progress_val = tk.DoubleVar(value=0)

        self.include_table = tk.BooleanVar(value=True)
        self.include_chart = tk.BooleanVar(value=True)
        self.include_stats = tk.BooleanVar(value=True)

        self._setup_ui()

    def _setup_ui(self):
        # Header
        header = tk.Frame(self.root, bg="#2c3e50", pady=18)
        header.pack(fill=tk.X)
        tk.Label(header, text="Excel  →  Word 보고서 생성기",
                 font=("맑은 고딕", 16, "bold"), fg="white", bg="#2c3e50").pack()
        tk.Label(header, text="엑셀 파일을 선택하면 자동으로 Word 보고서를 만들어 드립니다",
                 font=("맑은 고딕", 9), fg="#bdc3c7", bg="#2c3e50").pack()

        body = tk.Frame(self.root, bg="#f0f2f5", padx=25, pady=20)
        body.pack(fill=tk.BOTH, expand=True)

        def row_label(parent, text, r):
            tk.Label(parent, text=text, font=("맑은 고딕", 10),
                     bg="#f0f2f5", anchor="w").grid(row=r, column=0, sticky="w", pady=6)

        # Report title
        row_label(body, "보고서 제목", 0)
        tk.Entry(body, textvariable=self.report_title, font=("맑은 고딕", 10),
                 width=45, relief="solid", bd=1).grid(row=0, column=1, columnspan=2, sticky="ew", pady=6, padx=(12, 0))

        # Excel file
        row_label(body, "엑셀 파일", 1)
        tk.Entry(body, textvariable=self.excel_path, font=("맑은 고딕", 10),
                 width=38, state="readonly", relief="solid", bd=1).grid(row=1, column=1, sticky="ew", pady=6, padx=(12, 6))
        tk.Button(body, text="찾아보기", command=self._browse_excel,
                  bg="#3498db", fg="white", font=("맑은 고딕", 9), relief="flat",
                  padx=8).grid(row=1, column=2, pady=6)

        # Output file
        row_label(body, "저장 위치", 2)
        tk.Entry(body, textvariable=self.output_path, font=("맑은 고딕", 10),
                 width=38, state="readonly", relief="solid", bd=1).grid(row=2, column=1, sticky="ew", pady=6, padx=(12, 6))
        tk.Button(body, text="저장 위치", command=self._browse_output,
                  bg="#3498db", fg="white", font=("맑은 고딕", 9), relief="flat",
                  padx=8).grid(row=2, column=2, pady=6)

        body.columnconfigure(1, weight=1)

        # Options
        opt_frame = tk.LabelFrame(body, text="  포함 항목  ", font=("맑은 고딕", 10),
                                  bg="#f0f2f5", bd=1, relief="solid", padx=15, pady=10)
        opt_frame.grid(row=3, column=0, columnspan=3, sticky="ew", pady=12)

        for i, (var, label) in enumerate([
            (self.include_table, "데이터 표"),
            (self.include_chart, "차트 / 그래프"),
            (self.include_stats, "통계 요약"),
        ]):
            tk.Checkbutton(opt_frame, text=label, variable=var,
                           font=("맑은 고딕", 10), bg="#f0f2f5",
                           activebackground="#f0f2f5").grid(row=0, column=i, padx=20)

        # Progress
        self.progress_bar = ttk.Progressbar(body, variable=self.progress_val, maximum=100, length=560)
        self.progress_bar.grid(row=4, column=0, columnspan=3, sticky="ew", pady=(10, 4))
        tk.Label(body, textvariable=self.status, font=("맑은 고딕", 9),
                 fg="#7f8c8d", bg="#f0f2f5").grid(row=5, column=0, columnspan=3)

        # Generate button
        self.gen_btn = tk.Button(body, text="▶  보고서 생성",
                                  command=self._start_generation,
                                  bg="#27ae60", fg="white",
                                  font=("맑은 고딕", 13, "bold"),
                                  relief="flat", pady=10, padx=30,
                                  activebackground="#219a52", activeforeground="white")
        self.gen_btn.grid(row=6, column=0, columnspan=3, pady=18)

    # ── File dialogs ──────────────────────────────────────────────────────────

    def _browse_excel(self):
        path = filedialog.askopenfilename(
            title="엑셀 파일 선택",
            filetypes=[("Excel 파일", "*.xlsx *.xls *.xlsm"), ("모든 파일", "*.*")]
        )
        if path:
            self.excel_path.set(path)
            base = os.path.splitext(path)[0]
            self.output_path.set(base + "_보고서.docx")
            self.status.set(f"선택됨: {os.path.basename(path)}")
            self.progress_val.set(0)

    def _browse_output(self):
        path = filedialog.asksaveasfilename(
            title="저장 위치 선택",
            defaultextension=".docx",
            filetypes=[("Word 파일", "*.docx"), ("모든 파일", "*.*")]
        )
        if path:
            self.output_path.set(path)

    # ── Generation ────────────────────────────────────────────────────────────

    def _start_generation(self):
        if not self.excel_path.get():
            messagebox.showerror("오류", "엑셀 파일을 선택해주세요.")
            return
        if not self.output_path.get():
            messagebox.showerror("오류", "저장 위치를 선택해주세요.")
            return

        self.gen_btn.config(state="disabled")
        t = threading.Thread(target=self._generate, daemon=True)
        t.start()

    def _update(self, msg, pct=None):
        self.status.set(msg)
        if pct is not None:
            self.progress_val.set(pct)
        self.root.update_idletasks()

    def _generate(self):
        try:
            self._update("엑셀 파일 읽는 중...", 5)
            xl = pd.ExcelFile(self.excel_path.get())
            sheets = xl.sheet_names

            doc = Document()
            self._set_default_font(doc)

            self._update("표지 생성 중...", 10)
            self._cover_page(doc)

            total = len(sheets)
            for idx, name in enumerate(sheets):
                pct = 15 + (idx / total) * 78
                self._update(f"시트 처리 중: {name}  ({idx + 1}/{total})", pct)

                df = pd.read_excel(self.excel_path.get(), sheet_name=name)
                if df.empty:
                    continue

                doc.add_heading(name, level=1)

                if self.include_stats.get():
                    self._add_stats(doc, df)
                if self.include_table.get():
                    self._add_table(doc, df)
                if self.include_chart.get():
                    self._add_chart(doc, df, name)

                if idx < total - 1:
                    doc.add_page_break()

            self._update("파일 저장 중...", 96)
            doc.save(self.output_path.get())
            self._update("완료!", 100)

            if messagebox.askyesno("완료", f"보고서 생성 완료!\n\n{self.output_path.get()}\n\n지금 파일을 열어볼까요?"):
                os.startfile(self.output_path.get())

        except Exception as exc:
            messagebox.showerror("오류", f"보고서 생성 실패:\n{exc}")
            self._update("오류 발생", 0)
        finally:
            self.gen_btn.config(state="normal")

    # ── Document builders ─────────────────────────────────────────────────────

    def _set_default_font(self, doc):
        style = doc.styles['Normal']
        style.font.name = '맑은 고딕'
        style.font.size = Pt(10)

    def _cover_page(self, doc):
        for _ in range(6):
            doc.add_paragraph()

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(self.report_title.get())
        r.bold = True
        r.font.size = Pt(30)
        r.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

        doc.add_paragraph()

        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(f"생성일: {datetime.now().strftime('%Y년 %m월 %d일')}")
        r2.font.size = Pt(12)
        r2.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r3 = p3.add_run(f"원본 파일: {os.path.basename(self.excel_path.get())}")
        r3.font.size = Pt(11)
        r3.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

        doc.add_page_break()

    def _add_stats(self, doc, df):
        num_cols = df.select_dtypes(include='number').columns.tolist()
        if not num_cols:
            return

        p = doc.add_paragraph()
        run = p.add_run("통계 요약")
        run.bold = True
        run.font.size = Pt(11)

        stats = df[num_cols].describe().round(2)
        tbl = doc.add_table(rows=len(stats) + 1, cols=len(num_cols) + 1)
        tbl.style = 'Light Shading Accent 1'

        tbl.cell(0, 0).text = "통계"
        for j, col in enumerate(num_cols):
            tbl.cell(0, j + 1).text = str(col)
        for i, row_idx in enumerate(stats.index):
            tbl.cell(i + 1, 0).text = str(row_idx)
            for j, col in enumerate(num_cols):
                tbl.cell(i + 1, j + 1).text = str(stats.loc[row_idx, col])

        doc.add_paragraph()

    def _add_table(self, doc, df):
        p = doc.add_paragraph()
        run = p.add_run("데이터 표")
        run.bold = True
        run.font.size = Pt(11)

        MAX_ROWS = 100
        display = df.head(MAX_ROWS)
        if len(df) > MAX_ROWS:
            note = doc.add_paragraph()
            note.add_run(f"※ 상위 {MAX_ROWS}행만 표시합니다 (전체 {len(df)}행)").italic = True

        tbl = doc.add_table(rows=len(display) + 1, cols=len(display.columns))
        tbl.style = 'Table Grid'

        # Header
        for j, col_name in enumerate(display.columns):
            cell = tbl.rows[0].cells[j]
            cell.text = str(col_name)
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), '2C3E50')
            tcPr.append(shd)

        # Rows
        for i, (_, row) in enumerate(display.iterrows()):
            for j, val in enumerate(row):
                tbl.cell(i + 1, j).text = "" if pd.isna(val) else str(val)

        doc.add_paragraph()

    def _add_chart(self, doc, df, sheet_name):
        num_cols = df.select_dtypes(include='number').columns.tolist()
        if not num_cols:
            return

        p = doc.add_paragraph()
        run = p.add_run("차트")
        run.bold = True
        run.font.size = Pt(11)

        chart_cols = num_cols[:4]
        n_charts = min(len(chart_cols), 2)

        fig, axes = plt.subplots(1, n_charts, figsize=(12, 5))
        if n_charts == 1:
            axes = [axes]

        COLORS = ['#3498db', '#e74c3c', '#2ecc71', '#f39c12']

        for ax_idx, col in enumerate(chart_cols[:n_charts]):
            ax = axes[ax_idx]
            data = df[col].dropna()

            if len(data) <= 30:
                ax.bar(range(len(data)), data.values,
                       color=COLORS[ax_idx], alpha=0.85, edgecolor='white')
                ax.set_title(col, fontsize=12, fontweight='bold')
                ax.set_xlabel("행 번호")
            else:
                ax.hist(data.values, bins=20,
                        color=COLORS[ax_idx], alpha=0.85, edgecolor='white')
                ax.set_title(f"{col} 분포", fontsize=12, fontweight='bold')
                ax.set_xlabel(col)
                ax.set_ylabel("빈도")

            ax.set_ylabel(col if len(data) <= 30 else "빈도")
            ax.grid(axis='y', alpha=0.3)
            ax.spines['top'].set_visible(False)
            ax.spines['right'].set_visible(False)

        fig.suptitle(f"[ {sheet_name} ] 차트", fontsize=14, fontweight='bold')
        plt.tight_layout()

        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        buf.seek(0)
        plt.close(fig)

        doc.add_picture(buf, width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()


def main():
    root = tk.Tk()
    ExcelToWordApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
